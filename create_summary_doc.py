from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Resumen_RutaControl_FTEX.docx'
doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.8)
sec.left_margin = sec.right_margin = Inches(0.85)

def font(run, size=11, color='20342C', bold=False):
    run.font.name = 'Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.bold = bold

def p(text='', size=11, color='20342C', bold=False, align=None, after=6):
    para = doc.add_paragraph(); para.paragraph_format.space_after = Pt(after); para.paragraph_format.line_spacing = 1.15
    if align: para.alignment = align
    font(para.add_run(text), size, color, bold)
    return para

def heading(text):
    para = doc.add_paragraph(); para.paragraph_format.space_before = Pt(12); para.paragraph_format.space_after = Pt(5)
    font(para.add_run(text), 14, '17633C', True); return para

def bullet(text):
    para = doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after = Pt(3); para.paragraph_format.line_spacing = 1.15
    font(para.add_run(text), 10.5)

title = p('RutaControl', 24, '17633C', True, WD_ALIGN_PARAGRAPH.CENTER, 2)
p('Resumen del proyecto de control vehicular', 13, '527D2F', False, WD_ALIGN_PARAGRAPH.CENTER, 3)
p('FRUTOS TROPICALES EXPORT. PERÚ', 10, 'C77B17', True, WD_ALIGN_PARAGRAPH.CENTER, 16)

heading('Objetivo')
p('Aplicación web para controlar las salidas, llegadas, ubicación GPS, mantenimiento y abastecimiento de combustible de la flota. El diseño busca que el registro sea rápido desde un teléfono y que el administrador conserve evidencias verificables.')

heading('Flujo de recorridos')
bullet('Salida rápida: foto de placa para reconocer el vehículo, fecha y hora automáticas, GPS obligatorio para el origen y foto del odómetro inicial.')
bullet('Llegada: selección del vehículo en ruta, GPS para registrar el destino real y foto del odómetro final.')
bullet('El sistema calcula los kilómetros recorridos a partir del odómetro final menos el inicial.')
bullet('El botón de confirmación se bloquea hasta completar los datos obligatorios.')
bullet('Se puede adjuntar evidencia opcional del estado del vehículo al salir: estado, foto o video corto y observación breve.')

heading('Mapa y ubicación')
bullet('El GPS comienza automáticamente al confirmar una salida y se detiene al confirmar la llegada.')
bullet('El mapa del Inicio permite visualizar el trayecto registrado mientras la aplicación permanece abierta.')
bullet('El origen y destino se guardan con ubicación GPS; se intenta mostrar la dirección del origen automáticamente.')

heading('Mantenimiento y afinamiento')
bullet('Foto de placa para reconocer el vehículo y foto de odómetro para completar kilometraje.')
bullet('Registro técnico: tipo de servicio, fecha del servicio, próxima fecha y próximo kilometraje.')
bullet('Interfaz compacta con identidad visual inspirada en mango y Frutos Tropicales.')

heading('Combustible')
bullet('Fecha y hora se registran automáticamente al guardar.')
bullet('Foto del odómetro para completar kilometraje.')
bullet('Foto del comprobante para intentar completar el grifo/proveedor y el monto final o costo total.')
bullet('El campo de litros se eliminó para mantener el flujo simple.')

heading('Interfaz y tecnología')
bullet('Proyecto migrado a React con Vite; los datos se conservan actualmente en el navegador local.')
bullet('Diseño responsive con estilo verde/mango, botones de Salida y Llegada y formularios con evidencia fotográfica.')
bullet('Reconocimiento OCR local para leer placas, odómetros y comprobantes, sujeto a revisión del usuario.')
bullet('Mapa con OpenStreetMap y seguimiento GPS desde el navegador.')
bullet('Preparado como PWA instalable: ícono, manifiesto y soporte básico sin conexión.')

heading('Pendientes para producción')
bullet('Publicar la aplicación con HTTPS para que se instale en móviles y para permisos confiables de GPS/cámara.')
bullet('Conectar Supabase para cuentas de administrador y chofer, datos compartidos, fotos/videos y seguridad.')
bullet('Crear login de chofer para completar su nombre automáticamente y limitarle el acceso solo a sus registros.')
bullet('Para seguimiento GPS con el celular bloqueado, evaluar una app móvil nativa o una PWA avanzada.')

footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(footer.add_run('RutaControl · Frutos Tropicales Export. Perú'), 9, '738078')
doc.save(OUT)
